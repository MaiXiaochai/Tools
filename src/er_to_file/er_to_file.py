# -*- coding: utf-8 -*-

# @File:     er_to_file.py
# @Project:  Tools
# @Date:     2018/10/11 20:23
# @Author:   MaiXiaochai

from docx import Document
from docx.shared import Pt, RGBColor

from orcl_pools import OrclPool


class ErTransUtils(object):
    """
    从oracle读取指定表，提取指定格式的ER，保存到文件。
    解放劳动力，提高效率。将更多的时间投入到更有意义的事情上。
    """

    def __init__(self, file_path):
        self.doc = Document()
        self.save_path = file_path
        
    @staticmethod
    def deal_line(row_data):
        """
        处理一行数据为标准格式，注意只是返回数据，不返回表头
        :param row_data:    tuple   待处理的一行数据
        :return:            list    处理好的一行数据
        
        demo:
        -------------------------------[ 处理前 ]-----------------------------------------
        column_id, column_name, data_type, data_length, data_precision, nullable, comments
        (1,         'xxxTNUM', 'VARCHAR2', 120,        None,           'N',       None)
        (2,         'xxxENT',  'VARCHAR2', 120,        None,           'Y',       None)
        ----------------------------------------------------------------------------------

        -------------------------[ 处理后 ]------------------------
        序号	        字段名	    字段类型	    字段释义	        备注
        1		    ID	        NUMBER	    主键	自定义ID，  按1自增
        ----------------------------------------------------------
        """

        # 序号          字段名称
        # column_id     column_name
        res_line = list(row_data[:2])

        # 字段类型
        # data_type(data_length[, data precision])
        if not row_data[3] or "TIMESTAMP" in row_data[2].upper():
            res_line.append(row_data[2].upper())

        else:
            len_prcs_lis = [str(row_data[3]), str(row_data[4])] if row_data[4] else [str(row_data[3]), ]
            len_prcs = "{}({})".format(row_data[2], ','.join(len_prcs_lis))
            res_line.append(len_prcs)

        # 字段释义
        # comments
        desc = row_data[-1] if row_data[-1] else ''
        res_line.append(desc)

        # 备注
        # nullable
        # 这里的备注暂时用nullable值,若需要添加其他字段，再增加
        null_val = 'NOT NULL' if row_data[5].upper() == 'N' else ''

        # NOT NULL
        remark = null_val
        res_line.append(remark)

        return res_line

    def to_word(self, tbl_name, head_row, rows):
        """
        保存到world文档,基于docx库
        :param tbl_name:    str             表名称
        :param head_row:    list            表格标题行
        :param rows:        list/tuple      所有处理好的表字段描述
        :return:
        """

        # 添加表名称
        prgh = self.doc.add_paragraph()
        run = prgh.add_run(u'{}'.format(tbl_name))

        # 设置字体格式
        run.font.name = u'宋体(中文正文)'

        # 三号 16pt
        run.font.size = Pt(16)

        # 黑色
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # 粗体
        run.font.bold = True

        # 添加表格
        row_len, col_len = len(rows) + 1, len(head_row)
        tbl = self.doc.add_table(rows=row_len, cols=col_len, style='Table Grid')

        # 设置表格字体
        tbl.style.font.name = u'宋体(中文正文)'

        # 五号 10.5pt, 小五 9pt
        tbl.style.font.size = Pt(10.5)
        # tbl.style.

        # 填写表头
        ErTransUtils.__fill_word_tbl(tbl, 0, head_row)

        # 填写具体内容
        for i in range(1, row_len):
            ErTransUtils.__fill_word_tbl(tbl, i, rows[i - 1])

        # 添加一行空行
        self.doc.add_paragraph('\n')

    @staticmethod
    def __fill_word_tbl(tbl_obj, row_idx, row_data):
        """
        很小的功能，填写word表格的一行数据,基于docx
        :param tbl_obj:     doc.table.Table     word中表格对象
        :param row_idx:     int                 行序号
        :param row_data:    list                行数据
        :return:
        """

        tbl_cells = tbl_obj.rows[row_idx].cells

        for i in range(len(row_data)):
            tbl_cells[i].text = str(row_data[i])

    @staticmethod
    def fetch_tbl_names(orcl_obj):
        """
        获取当前oracle用户下的所有表名称
        :param orcl_obj:    object      OrclPool实例
        :return:            list       表名称
        """

        fetch_sql = "SELECT distinct table_name FROM user_tables"
        all_data = orcl_obj.fetch_all(fetch_sql)
        return [x[0] for x in all_data]

    def __del__(self):
        self.doc.save(self.save_path)


def er_to_file(db_cfg, fetch_sql, tbl_names, head_row, file_path):
    """
    保存oracle中指定的ER到doc文件。
    :param db_cfg:      dict    连接oracle的配置信息
    :param fetch_sql:   str     查询ER结构的语句
    :param tbl_names:   list    要导出ER表的表名称
    :param head_row:    list    表格标题行内容
    :param file_path:   str     ER保存到的xlsx文件
    :return:
    """

    orcl = OrclPool(db_cfg)
    doc = ErTransUtils(file_path)

    # 若tbl_names未指定或者说为空，则获取该用户下所有表名称
    tbl_names = tbl_names if tbl_names else ErTransUtils.fetch_tbl_names(orcl)

    count = 1
    for tbl_name in tbl_names:
        tbl_name = tbl_name.upper()

        print('[ NO.{}: | {} | processing ... ]'.format(count, tbl_name))
        tbl_info = orcl.fetch_all(fetch_sql.format(tbl_name))

        tmp_rows = []
        for i in tbl_info:
            tmp_rows.append(ErTransUtils.deal_line(i))

        # 编辑doc对象
        doc.to_word(tbl_name, head_row, tmp_rows)
        print('[ NO.{}: | {} | done. ]'.format(count, tbl_name))

        count += 1
    print('[ All work done. ]')


def main():
    """
    一些配置
    :return:
    """

    # ------------------------------[ 谨慎修改 ]----------------------------------
    # ER结构查询语句
    # 若修改该语句，请一并修改与其相关的其他逻辑和设置
    fetch_sql = """
                    SELECT
                        t.column_id,
                        t.column_name,
                        t.data_type,
                        t.data_length,
                        t.data_precision,
                        t.nullable,
                        c.comments
                    FROM
                        user_tab_columns t,user_col_comments c
                    WHERE
                        t .table_name = '{}'
                    AND t.Table_Name = c.table_name
                    AND t.COLUMN_name = c.column_name
                    ORDER BY
                        t.column_id
                    """

    # ER表格的表头
    head_row = [
        '序号',
        '字段名',
        '字段类型',
        '字段释义',
        '备注'
    ]

    # 导出文件路径
    path = './ER.doc'
    # ----------------------------------------------------------------------------

    # oracle数据库连接配置
    orcl_cfg = {
        'user': 'passwd',
        'passwd': 'passwd',
        'host': 'ip_str',
        'port': 1521,
        'sid': 'sid'
    }

    # 需要导出ER到文件的表的名称
    tbl_names = [
        'table_name1',
        'table_name2',
        'table_name3'
    ]

    er_to_file(orcl_cfg, fetch_sql, tbl_names, head_row, path)


if __name__ == '__main__':
    main()
